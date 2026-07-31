from docx import Document
import re
from config import METRICAS, PALABRAS_POR_PAGINA


# ==============================
#        CONFIGURACIÓN
# ==============================


NIVELES_ESTILO = {

    "Heading 1": 0,
    "Heading 2": 1,
    "Heading 3": 2

}



def abrir_documento(ruta):
    
    return Document(ruta)

# ==============================
#  ANALISIS GRAL DEL DOCUMENTO
# ==============================


def analizar_documento(ruta):

    documento = abrir_documento(ruta)

    texto = ""

    for parrafo in documento.paragraphs:
        texto += parrafo.text + "\n"

    palabras = len(texto.split())
    parrafos = len(documento.paragraphs)
    tablas = len(documento.tables)
    figuras = len(documento.inline_shapes)
    
    resultado = {}
    
    for metrica in METRICAS:
        
        resultado[metrica] = 0
        
    resultado["palabras"] = palabras
    resultado["parrafos"] = parrafos
    resultado["tablas"] = tablas
    resultado["figuras"] = figuras
    resultado["paginas"] = max(
        1,
        round(resultado["palabras"] / PALABRAS_POR_PAGINA)
    )
    resultado["referencias"] = 0 # para mas adelante

    return resultado


# ==============================
#    DETECCIÓN DE ESTRUCTURA
# ==============================

def obtener_nivel_estilo(estilo):

    if estilo.name in NIVELES_ESTILO:

        return NIVELES_ESTILO[estilo.name]


    if estilo.base_style:

        if estilo.base_style.name in NIVELES_ESTILO:

            return NIVELES_ESTILO[estilo.base_style.name]


    return None


def generar_id_seccion(titulo, nivel, contadores):

    titulo = titulo.strip()

    coincidencia = None


    # Caso 1: numeración al inicio
    coincidencia = re.match(
        r"^(\d+(?:\.\d+)*)",
        titulo
    )


    if coincidencia:

        numeros = coincidencia.group(1).replace(".", "_")

        if nivel == 0:
            return f"CAP_{numeros}"

        elif nivel == 1:
            return f"SEC_{numeros}"

        elif nivel == 2:
            return f"SUB_{numeros}"


    # Caso 2: palabras tipo Capítulo / Chapter
    if nivel == 0:

        coincidencia = re.search(
            r"(?:cap[ií]tulo|chapter|cap[ií]tulo)\s+(\d+)",
            titulo,
            re.IGNORECASE
        )

        if coincidencia:

            numero = coincidencia.group(1)

            return f"CAP_{numero}"


    # Caso 3: sin numeración
    contadores[nivel] += 1


    if nivel == 0:

        return f"CAP_{contadores[0]:03d}"


    elif nivel == 1:

        return f"SEC_{contadores[1]:03d}"


    elif nivel == 2:

        return f"SUB_{contadores[2]:03d}"



def obtener_secciones(ruta):

    documento = abrir_documento(ruta)

    secciones = []
    
    contadores = {
        
        0: 0,
        1: 0,
        2:0
        
    }
    

    for parrafo in documento.paragraphs:

        nivel = obtener_nivel_estilo(parrafo.style)

        if nivel is not None and parrafo.text.strip():
            
            titulo = parrafo.text.strip()
            
            identificador = generar_id_seccion(
                titulo,
                nivel,
                contadores
            )

            secciones.append({
                "id": identificador,
                "titulo": parrafo.text,
                "nivel": nivel
            })

    return secciones



def construir_estructura(secciones):

    estructura = []

    capitulo_actual = None
    seccion_actual = None


    for item in secciones:


        if item["nivel"] == 0:

            capitulo_actual = {
                "id": item["id"],
                "titulo": item["titulo"],
                "nivel": 0,
                "secciones": []
            }

            estructura.append(capitulo_actual)

            seccion_actual = None


        elif item["nivel"] == 1:

            if capitulo_actual is not None:

                seccion_actual = {
                    "id": item["id"],
                    "titulo": item["titulo"],
                    "nivel": 1,
                    "subsecciones": []
                }

                capitulo_actual["secciones"].append(
                    seccion_actual
                )


        elif item["nivel"] == 2:

            if seccion_actual is not None:

                seccion_actual["subsecciones"].append({
                    "id": item["id"],
                    "titulo": item["titulo"],
                    "nivel": 2
                })


    return estructura


def mostrar_estructura(estructura):

    for capitulo in estructura:

        print()
        print(
            capitulo["id"],
            "-",
            capitulo["titulo"]
        )


        for seccion in capitulo["secciones"]:

            print(
                "   └──",
                seccion["id"],
                "-",
                seccion["titulo"]
            )


            for subseccion in seccion["subsecciones"]:

                print(
                    "          └──",
                    subseccion["id"],
                    "-",
                    subseccion["titulo"]
                )
                

# ==============================
#     DETECCIÓN DE ELEMENTOS
# ==============================

def obtener_elementos_documento(documento):

    cuerpo = documento.element.body

    for elemento in cuerpo:

        if elemento.tag.endswith("}p"):

            yield "parrafo", elemento


        elif elemento.tag.endswith("}tbl"):

            yield "tabla", elemento
            

def tiene_figura(parrafo):

    return "graphic" in parrafo._p.xml



def tiene_tabla(elemento):
    
    return elemento.tag.endswith("tbl")


# ==============================
#    ANÁLISIS POR SECCIONES
# ==============================

def analizar_secciones(ruta):

    documento = abrir_documento(ruta)

    secciones = []

    seccion_actual = None

    contadores = {
        0: 0,
        1: 0,
        2: 0
    }


    for parrafo in documento.paragraphs:

        nivel = obtener_nivel_estilo(parrafo.style)


        if nivel is not None and parrafo.text.strip():

            titulo = parrafo.text.strip()


            identificador = generar_id_seccion(
                titulo,
                nivel,
                contadores
            )


            seccion_actual = {

                "id": identificador,
                "titulo": titulo,
                "nivel": nivel,
                "metricas": {}

            }


            for metrica in METRICAS:

                seccion_actual["metricas"][metrica] = 0


            secciones.append(seccion_actual)


        else:

            if seccion_actual is not None:

                palabras = len(parrafo.text.split())

                seccion_actual["metricas"]["palabras"] += palabras
                
                seccion_actual["metricas"]["parrafos"] += 1


    return secciones



def analizar_metricas_secciones(ruta, secciones):

    documento = abrir_documento(ruta)

    seccion_actual = None


    # creamos un diccionario para encontrar rápido los párrafos
    parrafos = {}

    for parrafo in documento.paragraphs:
        parrafos[parrafo._p] = parrafo



    for elemento in documento.element.body:


        # -----------------------
        # TABLAS
        # -----------------------

        if tiene_tabla(elemento):

            if seccion_actual is not None:

                seccion_actual["metricas"]["tablas"] += 1



        # -----------------------
        # PÁRRAFOS
        # -----------------------

        elif elemento.tag.endswith("p"):


            parrafo = parrafos.get(elemento)


            if parrafo is None:
                continue


            texto = parrafo.text.strip()



            # primero revisamos figura
            if tiene_figura(parrafo):

                if seccion_actual is not None:

                    seccion_actual["metricas"]["figuras"] += 1



            # después vemos si es título

            if texto:


                nivel = obtener_nivel_estilo(parrafo.style)


                if nivel is not None:


                    for seccion in secciones:

                        if seccion["titulo"].strip() == texto:

                            seccion_actual = seccion
                            break



                else:


                    if seccion_actual is not None:

                        seccion_actual["metricas"]["palabras"] += len(
                            texto.split()
                        )


    return secciones

